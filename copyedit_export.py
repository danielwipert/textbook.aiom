#!/usr/bin/env python3
"""
copyedit_export.py  |  chapter HTML to a copy-editing .docx, plus a round-trip manifest

Stage 6 is Dan's, run outside the Claude system, and Word is where it happens.
This pass turns the live chapter HTML into a document that can be copy edited
line by line, and writes a manifest so the returned file can be mapped back into
the HTML block by block rather than by eye.

What travels, and why:

  Every editable block goes, in document order, one HTML block to one Word
  paragraph. Each carries a bracketed type tag, so the editor can see at a
  glance whether a line is body prose, a definition, a footnote gloss, or a
  problem label, and so the returned file can be verified against the type
  sequence before a single edit is imported.

  Figures travel as images cropped from the production PDF, with their captions
  and, separately, every text string set inside the figure. Those strings are
  prose too: nothing else in the pipeline puts them in front of a copy editor.

  The source register does NOT travel. It is audit apparatus under Decision 51,
  it is machine-read, and a copy edit inside it would silently change what a
  citation claims. It is listed in the manifest as excluded so the omission is
  a recorded decision rather than an oversight.

Usage:
    python3 copyedit_export.py <chapter.html> --pdf <production.pdf> --out <name>

Writes <name>.docx and <name>.manifest.json. Import with copyedit_import.py.
"""
import argparse
import hashlib
import html as htmllib
import json
import os
import re
import sys

# The key-term register sets its term names in a <span>, not a <p>, so a
# paragraph-only sweep drops all seven. They are prose: they must match the
# definition callouts word for word, and nothing else in the pipeline puts them
# in front of a copy editor. Caught by the completeness audit, 2026-08-06.
# <li> added 2026-08-08. It was absent, so THM-009's four antecedents, which are
# <li> elements inside <ol class="ante">, were in no proof ever issued. The
# Chapter 1 round-1 return rebuilt the theorem from the older running-prose
# paraphrase because that is all the editor could see, and one registry
# antecedent was dropped and one added. The most rule-bound prose in the book was
# the only prose Stage 6 could not show a human.
BLOCK_RE = re.compile(
    r'(?s)<(p|h1|h2|h3|figcaption|td|th|li|span)\b([^>]*)>(.*?)</\1>')

# Spans are used inline throughout body prose; only these carry a block.
SPAN_BLOCK_CLASSES = {"kt-t"}

# A theorem panel is a RENDERING of a locked registry statement, not prose
# (standing rule 6, Decision 56). These kinds travel so they can be proofread and
# are labelled so they are not rewritten; copyedit_import.py refuses to apply an
# edit to any of them.
REGISTRY_BOUND = {"THEOREM STATEMENT", "THEOREM SCOPE",
                  "THEOREM ANTECEDENT", "THEOREM CONSEQUENT"}

# Class to tag. The tag is what the editor reads and what the importer keys on.
TAGS = {
    "part-label": "PART LABEL",
    "slot-label": "SLOT",
    "case-title": "CASE TITLE",
    "provenance": "PROVENANCE",
    "lab": "LABEL",
    "term": "TERM",
    "stmt": "THEOREM STATEMENT",
    "cond": "THEOREM SCOPE",
    "conseq": "THEOREM CONSEQUENT",
    "date": "DATED",
    "mlab": "SUBHEAD",
    "kt-t": "KEY TERM",
    "plab": "PROBLEM LABEL",
    "ptitle": "PROBLEM TITLE",
    "pnote": "PROBLEM NOTE",
    "dq-n": "QUESTION NUMBER",
    "num": "SECTION NUMBER",
}


# Emphasis tags carry no width, so replacing one with a space invents a space
# the page does not have. Everything else, including span, is replaced with a
# space: span.num, span.fignum and span.mk all sit hard against the text that
# follows them, and "1.1The purchase" would be worse than the artifact this
# fixes. Added 2026-08-08, see strip() below.
INLINE_TAGS = ("b", "i", "em", "strong", "sup", "sub", "code", "abbr")
INLINE_RE = re.compile(r"(?s)</?(?:%s)\b[^>]*>" % "|".join(INLINE_TAGS))

# span.nb is the SAME case as the emphasis tags and was missed when they were
# fixed. It wraps a proper noun inside running prose (Decision 58), so it
# carries no width: spacing it turns "<span class="nb">Uber</span>\u2019s" into
# "Uber \u2019s" and "<span class="nb">Copilot</span>." into "Copilot .". The
# other span classes named above are prefix markers that sit hard against the
# text FOLLOWING them, which is why the generic rule still spaces them. Chapter
# 2 tripped this in two blocks of the Stage 6 proof, and the unedited round trip
# cannot see it for the reason strip() records: the artifact is symmetric.
NB_RE = re.compile(r'<span class="nb">([^<]*)</span>')


def strip(fragment):
    """Visible text of an HTML fragment, entities resolved, spacing normalised.

    Emphasis tags are removed rather than spaced. Replacing every tag with a
    space turned "<b>access price</b>." into "access price ." and put that
    phantom space in the copy-editing proof, where an editor correctly closed it
    up and the importer then refused the edit as unlocatable, because the space
    was never in the HTML. Three of Chapter 1's blocks did this in the round-2
    proof. The unedited round trip cannot see it: the artifact is symmetric, so
    export and import agree with each other and both differ from the page.
    """
    t = INLINE_RE.sub("", fragment)
    t = NB_RE.sub(r"\1", t)
    t = re.sub(r"(?s)<[^>]+>", " ", t)
    return re.sub(r"\s+", " ", htmllib.unescape(t)).strip()


def figure_text(svg):
    """Every text string set inside a figure, in document order.

    These are the only words in the chapter that no other pass shows a human:
    voicecheck strips SVG, the gates read the PDF for geometry and fonts, and
    the fact checks read prose. A typo here would ship.
    """
    return [strip(m.group(1))
            for m in re.finditer(r"(?s)<text[^>]*>(.*?)</text>", svg)
            if strip(m.group(1))]


def crop_figures(pdf_path, out_dir, count):
    """Crop each figure from the production PDF, between its lead-in and caption.

    Rendering the SVG directly would need a converter this environment does not
    carry, and would also show the figure as authored rather than as printed.
    The point of a copy-editing proof is what the reader will see.
    """
    import pdfplumber
    from pdf2image import convert_from_path

    pdf = pdfplumber.open(pdf_path)
    caps = []
    for i, page in enumerate(pdf.pages):
        for k, chars in _lines(page):
            txt = "".join(c["text"] for c in chars)
            if txt.startswith("Figure") and round(chars[0]["size"], 1) < 10.5:
                caps.append((i, k, chars))
    caps = caps[:count]

    paths = []
    for n, (pageno, cap_top, _) in enumerate(caps, start=1):
        page = pdf.pages[pageno]
        above = [k for k, chars in _lines(page)
                 if k < cap_top - 12 and round(chars[0]["size"], 1) >= 10.5]
        top = (max(above) + 14) if above else 60
        img = convert_from_path(pdf_path, dpi=200,
                                first_page=pageno + 1, last_page=pageno + 1)[0]
        scale = img.height / float(page.height)
        box = (0, int(top * scale), img.width, int((cap_top - 3) * scale))
        path = os.path.join(out_dir, f"figure_{n}.png")
        img.crop(box).save(path)
        paths.append(path)
    return paths


def _lines(page):
    rows = {}
    for c in page.chars:
        if c["text"].strip():
            rows.setdefault(round(c["top"], 0), []).append(c)
    return [(k, rows[k]) for k in sorted(rows)]


def extract(html_path):
    """Ordered editable blocks, each carrying its exact span in the source file.

    The span is the point. An importer that searches the whole file for a
    changed phrase finds "June 16 and" three times, twice inside the source
    register, and has to refuse. Anchored to its own block it finds it once and
    applies it. Everything here therefore works on the ORIGINAL text, with no
    rewriting beforehand, so every offset recorded is an offset into the file
    that will later be edited.

    The source register is excluded by cutting the body at its opening tag, so
    no span can ever point inside it.
    """
    src = open(html_path, encoding="utf-8").read()
    cut = src.find('<section id="aiom-sources">')
    body_end = cut if cut >= 0 else len(src)

    figures = [(m.start(), m.end(), m.group(0))
               for m in re.finditer(r"(?s)<figure>.*?</figure>", src[:body_end])]
    fig_strings = [figure_text(f[2]) for f in figures]

    # A theorem antecedent is an <li> inside <ol class="ante">, and it opens with
    # <span class="mk">(i)</span>. The marker is structure, not prose: the span
    # recorded below starts AFTER it, so an edit can never renumber an antecedent.
    ante = [(m.start(), m.end()) for m in
            re.finditer(r'(?s)<ol class="ante">.*?</ol>', src[:body_end])]

    blocks = []
    fig_n = 0
    for m in BLOCK_RE.finditer(src[:body_end]):
        tag, attrs, inner = m.group(1), m.group(2), m.group(3)
        cls = re.search(r'class="([^"]*)"', attrs)
        cls = cls.group(1).strip() if cls else ""
        if "<svg" in inner:
            continue
        if tag == "span" and cls not in SPAN_BLOCK_CLASSES:
            continue

        inner_start = m.start(3)
        marker = None
        if tag == "li":
            mk = re.match(r'(?s)\s*<span class="mk">(.*?)</span>', inner)
            if mk:
                # Parentheses are stripped: the importer reads a leading
                # "(...)" on a tag line as the sources group, so a marker of
                # "(i)" would truncate the note at the first bracket.
                marker = strip(mk.group(1)).strip("()")
                inner = inner[mk.end():]
                inner_start += mk.end()
        cites = list(re.finditer(r"(?s)<cite\b[^>]*>(.*?)</cite>", inner))
        text = strip(re.sub(r"(?s)<cite\b[^>]*>.*?</cite>", " ", inner))

        if tag == "h1":
            kind = "CHAPTER TITLE"
        elif tag == "h2":
            kind = TAGS.get(cls, "SECTION TITLE")
        elif tag == "h3":
            kind = "SECTION HEAD"
        elif tag == "figcaption":
            fig_n += 1
            kind = "FIGURE CAPTION"
        elif tag == "th":
            kind = "TABLE HEADING"
        elif tag == "td":
            kind = "TABLE CELL"
        elif tag == "li":
            kind = ("THEOREM ANTECEDENT"
                    if any(a <= m.start() < b for a, b in ante) else "LIST ITEM")
        else:
            kind = TAGS.get(cls, "BODY")

        if text:
            blk = {
                "kind": kind, "text": text, "class": cls, "tag": tag,
                "figure": fig_n if kind == "FIGURE CAPTION" else None,
                "span": [inner_start, m.end(3)],
            }
            if kind in REGISTRY_BOUND:
                blk["sources"] = ((marker + ", " if marker else "")
                                  + "registry verbatim, do not edit")
            blocks.append(blk)

        for c in cites:
            gloss = re.sub(r'(?s)<span class="ckey">.*?</span>', " ", c.group(1))
            keys = re.search(r'<span class="ckey">\[(.*?)\]</span>', c.group(1))
            blocks.append({
                "kind": "FOOTNOTE", "text": strip(gloss), "class": "cite",
                "tag": "cite", "sources": keys.group(1) if keys else "",
                "figure": None,
                "span": [inner_start + c.start(1), inner_start + c.end(1)],
            })

        # A figure's caption is followed in the document by every string set
        # inside the figure, so the manifest has to follow the same order or the
        # importer reports a structure change on a file nobody touched.
        if kind == "FIGURE CAPTION":
            f_start, f_end, f_html = figures[fig_n - 1]
            for s_i, t in enumerate(
                    re.finditer(r"(?s)<text[^>]*>(.*?)</text>", f_html)):
                if not strip(t.group(1)):
                    continue
                blocks.append({
                    "kind": "FIGURE TEXT", "text": strip(t.group(1)),
                    "class": "svg-text", "tag": "text", "figure": fig_n,
                    "string": s_i,
                    "span": [f_start + t.start(1), f_start + t.end(1)],
                })

    return blocks, fig_strings


def build(blocks, fig_strings, fig_images, out_docx, chapter_title, source_name):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, Inches, RGBColor

    doc = Document()
    for name, size in (("Normal", 11),):
        st = doc.styles[name]
        st.font.name = "Georgia"
        st.font.size = Pt(size)
    doc.styles["Normal"].paragraph_format.space_after = Pt(10)
    doc.styles["Normal"].paragraph_format.line_spacing = 1.4

    def tagged(kind, text, sources=None):
        p = doc.add_paragraph()
        run = p.add_run(f"[{kind}] ")
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x88, 0x77)
        if sources:
            s = p.add_run(f"({sources}) ")
            s.font.size = Pt(8)
            s.font.color.rgb = RGBColor(0x99, 0x88, 0x77)
        p.add_run(text)
        return p

    h = doc.add_paragraph()
    r = h.add_run(chapter_title)
    r.bold = True
    r.font.size = Pt(16)

    note = doc.add_paragraph()
    note.add_run(
        "Copy-editing proof. Generated from " + source_name + ". Edit the prose "
        "freely, in place or with track changes, either works. Two things to "
        "leave alone: the grey bracketed tags, which tell the importer what each "
        "block is, and the order of the blocks. Splitting or merging a paragraph "
        "is fine and will be detected. The chapter's source register is not in "
        "this file: it is machine-read apparatus, and an edit inside it would "
        "change what a citation claims without changing what the page says. "
        "Blocks marked 'registry verbatim, do not edit' are a theorem panel, "
        "which renders a locked registry statement rather than stating one. They "
        "travel so you can proofread them and read the chapter whole. Flag "
        "anything wrong in them in a comment; the importer will not apply an "
        "edit made inside one, because changing an antecedent changes the "
        "theorem."
    ).italic = True

    doc.add_paragraph()     # spacer between the front matter and the chapter

    fig_seen = 0
    for b in blocks:
        if b["kind"] == "FIGURE CAPTION":
            fig_seen += 1
            idx = fig_seen - 1
            if idx < len(fig_images) and os.path.exists(fig_images[idx]):
                doc.add_picture(fig_images[idx], width=Inches(5.6))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tagged(b["kind"], b["text"], b.get("sources"))

    doc.save(out_docx)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("html")
    ap.add_argument("--pdf", required=True)
    ap.add_argument("--out", required=True)
    ap.add_argument("--title", default="")
    a = ap.parse_args()

    blocks, fig_strings = extract(a.html)
    out_dir = os.path.dirname(os.path.abspath(a.out)) or "."
    os.makedirs(out_dir, exist_ok=True)
    images = crop_figures(a.pdf, out_dir, len(fig_strings))

    title = a.title or os.path.basename(a.html)
    build(blocks, fig_strings, images, a.out + ".docx", title,
          os.path.basename(a.html))

    manifest = {
        "source_html": a.html,
        "source_pdf": a.pdf,
        "block_count": len(blocks),
        "excluded": ["<section id=\"aiom-sources\">, the Decision 51 source "
                     "register: machine-read apparatus, deliberately not exposed "
                     "to a copy edit"],
        "figure_text_counts": [len(f) for f in fig_strings],
        "blocks": [{"i": i, "kind": b["kind"], "class": b["class"],
                    "tag": b["tag"], "sources": b.get("sources"),
                    "figure": b.get("figure"), "string": b.get("string"),
                    "span": b["span"],
                    "sha1": hashlib.sha1(b["text"].encode()).hexdigest(),
                    "text": b["text"]}
                   for i, b in enumerate(blocks)],
    }
    with open(a.out + ".manifest.json", "w", encoding="utf-8") as f:
        json.dump(manifest, f, indent=1, ensure_ascii=False)

    kinds = {}
    for b in blocks:
        kinds[b["kind"]] = kinds.get(b["kind"], 0) + 1
    print(f"{len(blocks)} blocks, {sum(len(f) for f in fig_strings)} figure "
          f"strings, {len(images)} figure image(s)")
    for k in sorted(kinds):
        print(f"   {k:<20} {kinds[k]}")
    print(f"wrote {a.out}.docx and {a.out}.manifest.json")


if __name__ == "__main__":
    sys.exit(main())
